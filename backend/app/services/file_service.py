"""File service — upload, text extraction, OCR."""

from __future__ import annotations

import base64
import hashlib
import json
import os
import requests as http_requests
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import fitz  # PyMuPDF
import magic
from docx import Document as DocxDocument
from sqlalchemy.orm import Session as DBSession

from app.config import settings
from app.encryption import decrypt_api_key
from app.models.source_file import SourceFile
from app.models.file_content_cache import FileContentCache
from app.models.user import User
from app.models.secret_key import SecretKey
from app.models.ai_request import AIRequest

OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_REFERER = "https://edugen.local"
OPENROUTER_TITLE = "EduGen Local"


ALLOWED_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/jpeg",
    "image/png",
}

MIME_TO_TYPE = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/jpeg": "image",
    "image/png": "image",
}


def detect_mime(file_bytes: bytes) -> str:
    """Detect MIME type of uploaded file."""
    return magic.from_buffer(file_bytes, mime=True)


def compute_file_hash(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of file bytes for deduplication."""
    return hashlib.sha256(file_bytes).hexdigest()


def validate_file(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """Validate file size and type. Returns (mime_type, file_type)."""
    if len(file_bytes) > settings.max_file_size_bytes:
        raise ValueError(f"File exceeds maximum size of {settings.MAX_FILE_SIZE_MB}MB")

    mime_type = detect_mime(file_bytes)
    if mime_type not in ALLOWED_MIMES:
        raise ValueError(f"File type '{mime_type}' is not supported. Allowed: PDF, DOCX, JPG, PNG")

    return mime_type, MIME_TO_TYPE[mime_type]


def save_file(file_bytes: bytes, subject_id: str, extension: str) -> tuple[str, str]:
    """Save file to disk. Returns (file_uuid, relative_path)."""
    file_uuid = str(uuid4())
    folder = Path(settings.DATA_DIR) / "subjects" / subject_id
    folder.mkdir(parents=True, exist_ok=True)

    filename = f"{file_uuid}.{extension}"
    filepath = folder / filename
    filepath.write_bytes(file_bytes)

    return file_uuid, str(filepath)


def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF file. Returns (text, page_count)."""
    doc = fitz.open(file_path)
    page_count = len(doc)
    texts = []
    for page in doc:
        text = page.get_text()
        texts.append(text)
    doc.close()
    return "\n\n".join(texts), page_count


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n\n".join(paragraphs)


def is_scanned_pdf(file_path: str, min_chars_per_page: int = 100) -> bool:
    """Check if a PDF has very little text per page (likely scanned/image-based)."""
    doc = fitz.open(file_path)
    page_count = len(doc)
    if page_count == 0:
        doc.close()
        return True
    total_chars = sum(len(page.get_text().strip()) for page in doc)
    doc.close()
    avg_chars_per_page = total_chars / page_count
    return avg_chars_per_page < min_chars_per_page


def pdf_pages_to_images(file_path: str, max_pages: int = 5) -> list[bytes]:
    """Convert PDF pages to PNG images for OCR."""
    doc = fitz.open(file_path)
    images = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        pix = page.get_pixmap(dpi=200)
        images.append(pix.tobytes("png"))
    doc.close()
    return images


def ocr_image_with_vision(
    image_bytes: bytes,
    api_key: str,
    model: str = "openai/gpt-5-mini",
) -> str:
    """Use OpenRouter (Vision-capable model) to OCR an image."""
    b64 = base64.b64encode(image_bytes).decode()

    resp = http_requests.post(
        OPENROUTER_API_URL,
        headers={
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": OPENROUTER_REFERER,
            "X-OpenRouter-Title": OPENROUTER_TITLE,
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Transcribe all textual, structural, and mathematical components from this image exactly. Output the text in the original language.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{b64}"},
                        },
                    ],
                }
            ],
        },
        timeout=120,
    )

    if resp.status_code != 200:
        raise RuntimeError(f"OpenRouter OCR error ({resp.status_code}): {resp.text[:300]}")

    return resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")


def generate_summary(text: str, api_key: str, model: str = "openai/gpt-5-mini") -> str:
    """Generate a 1-sentence summary of extracted text via OpenRouter."""
    resp = http_requests.post(
        OPENROUTER_API_URL,
        headers={
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": OPENROUTER_REFERER,
            "X-OpenRouter-Title": OPENROUTER_TITLE,
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "Provide a 1-sentence descriptive summary of this educational material in the same language as the material.",
                },
                {"role": "user", "content": text[:4000]},
            ],
        },
        timeout=60,
    )

    if resp.status_code != 200:
        raise RuntimeError(f"OpenRouter summary error ({resp.status_code}): {resp.text[:300]}")

    return resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")


def get_api_key_and_model(db: DBSession, user_id: str | None = None) -> tuple[str, str]:
    """Retrieve and decrypt the API key and the user's preferred model.

    Looks up the first active key from the secret_keys table.
    The model preference is read from the users.default_model column.
    """
    api_key = None
    model = "openai/gpt-5-mini"

    if user_id:
        secret_key = (
            db.query(SecretKey)
            .filter(SecretKey.user_id == user_id, SecretKey.is_active == True)
            .first()
        )
        if secret_key:
            api_key = decrypt_api_key(secret_key.secret_key_hash)

        # Get model preference from the user record
        user = db.query(User).filter(User.id == user_id).first()
        if user and user.default_model:
            model = user.default_model

    if not api_key:
        raise ValueError("API key not configured")

    return api_key, model


def process_file_extraction(db: DBSession, source_file_id: str) -> None:
    """Background task: extract text from a file and generate summary.

    Checks the global FileContentCache first (keyed on SHA-256 of file bytes).
    If cached data is found the expensive AI calls are skipped entirely.
    """
    source_file = db.query(SourceFile).filter(SourceFile.id == source_file_id).first()
    if not source_file:
        return

    # --- Cache hit: reuse already-extracted content --------------------------
    if source_file.file_hash:
        cache_entry = db.query(FileContentCache).filter(
            FileContentCache.file_hash == source_file.file_hash
        ).first()
        if cache_entry:
            source_file.extracted_text = cache_entry.extracted_text
            source_file.summary = cache_entry.summary
            source_file.page_count = cache_entry.page_count
            db.commit()
            return
    # -------------------------------------------------------------------------

    try:
        api_key, model = get_api_key_and_model(db, user_id=source_file.user_id)
    except ValueError:
        api_key = None
        model = "openai/gpt-5-mini"

    file_path = source_file.original_path
    extracted_text = ""
    page_count = None

    try:
        if source_file.file_type == "pdf":
            extracted_text, page_count = extract_text_from_pdf(file_path)
            source_file.page_count = page_count

            # Check if scanned PDF
            if is_scanned_pdf(file_path):
                if api_key:
                    images = pdf_pages_to_images(file_path, max_pages=5)
                    ocr_texts = []
                    for img in images:
                        ocr_text = ocr_image_with_vision(img, api_key, model)
                        ocr_texts.append(ocr_text)
                    extracted_text = "\n\n".join(ocr_texts)
                else:
                    extracted_text = "[OCR_ERROR:NO_API_KEY]"

        elif source_file.file_type == "docx":
            extracted_text = extract_text_from_docx(file_path)

        elif source_file.file_type == "image":
            if not api_key:
                extracted_text = "[OCR_ERROR:NO_API_KEY]"
            else:
                image_bytes = Path(file_path).read_bytes()
                extracted_text = ocr_image_with_vision(image_bytes, api_key, model)

        source_file.extracted_text = extracted_text

        # Generate summary if we have text and API key
        summary = None
        if extracted_text.strip() and not extracted_text.startswith("[OCR_ERROR:") and api_key:
            summary = generate_summary(extracted_text, api_key, model)
            source_file.summary = summary

        db.commit()

        # --- Populate cache so future uploads of the same file are instant ---
        if source_file.file_hash and not extracted_text.startswith("[OCR_ERROR:"):
            try:
                existing = db.query(FileContentCache).filter(
                    FileContentCache.file_hash == source_file.file_hash
                ).first()
                if not existing:
                    cache_entry = FileContentCache(
                        file_hash=source_file.file_hash,
                        file_type=source_file.file_type,
                        extracted_text=extracted_text,
                        summary=summary,
                        page_count=page_count,
                        created_at=datetime.now(timezone.utc).isoformat(),
                    )
                    db.add(cache_entry)
                    db.commit()
            except Exception:
                # Cache write failure is non-critical
                db.rollback()
        # ---------------------------------------------------------------------

    except Exception as e:
        # Log error but don't crash
        source_file.extracted_text = f"[Extraction error: {str(e)}]"
        db.commit()
