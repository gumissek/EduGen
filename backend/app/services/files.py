from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile


def detect_file_type(upload: UploadFile) -> str:
    content_type = (upload.content_type or "").lower()
    if "pdf" in content_type:
        return "pdf"
    if "word" in content_type or "officedocument.wordprocessingml" in content_type:
        return "docx"
    if "image" in content_type:
        return "img"
    suffix = (upload.filename or "").split(".")[-1].lower()
    if suffix in {"pdf", "docx", "png", "jpg", "jpeg"}:
        return "img" if suffix in {"png", "jpg", "jpeg"} else suffix
    raise HTTPException(status_code=400, detail="Nieobsługiwany format pliku")


async def save_upload(upload: UploadFile, target_path: Path) -> int:
    data = await upload.read()
    if len(data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Plik przekracza 10MB")
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(data)
    return len(data)


def extract_pdf_text(path: Path) -> tuple[str, int]:
    text_chunks: list[str] = []
    page_count = 0
    with fitz.open(path) as doc:
        page_count = doc.page_count
        for page in doc:
            text_chunks.append(page.get_text("text"))
    return "\n".join(text_chunks).strip(), page_count


def extract_docx_text(path: Path) -> str:
    document = DocxDocument(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def extract_image_bytes(path: Path) -> bytes:
    return path.read_bytes()


def one_sentence_summary(text: str, fallback: str) -> str:
    cleaned = " ".join(text.split())
    if not cleaned:
        return f"Materiał źródłowy: {fallback}."
    short = cleaned[:180]
    if not short.endswith("."):
        short += "."
    return short


def build_docx_buffer(content: str, answer_key: str) -> io.BytesIO:
    file = io.BytesIO()
    document = DocxDocument()
    document.add_heading("EduGen - Materiał", level=1)
    for line in content.split("\n"):
        if line.strip():
            document.add_paragraph(line.strip())
    document.add_page_break()
    document.add_heading("Klucz odpowiedzi", level=2)
    for line in answer_key.split("\n"):
        if line.strip():
            document.add_paragraph(line.strip())
    document.save(file)
    file.seek(0)
    return file
