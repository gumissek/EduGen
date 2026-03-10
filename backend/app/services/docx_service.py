"""DOCX generation service with variant shuffling."""

from __future__ import annotations

import copy
import json
import random
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session as DBSession

from app.config import settings
from app.models.generation import Generation
from app.models.prototype import Prototype
from app.models.document import Document
from app.models.subject import Subject
from app.services.ai_service import TYPES_WITHOUT_QUESTIONS


CONTENT_TYPE_FOLDER_NAMES = {
    "worksheet": "Karta_pracy",
    "test": "Sprawdzian",
    "quiz": "Kartkowka",
    "exam": "Test",
    "lesson_materials": "Materialy_na_zajeciach",
}

EDUCATION_LEVEL_FOLDER_NAMES = {
    "primary": "Szkola_podstawowa",
    "secondary": "Szkola_srednia",
}


def _parse_content_to_questions(content: str) -> list[dict]:
    """Parse HTML content or JSON back to a list of question dicts.

    Tries JSON first, then falls back to a basic HTML parser.
    """
    # Try parsing as JSON first
    try:
        data = json.loads(content)
        if isinstance(data, dict) and "questions" in data:
            return data["questions"]
        if isinstance(data, list):
            return data
    except (json.JSONDecodeError, TypeError):
        pass

    # Simple HTML-based parser as fallback
    questions = []
    # Split by question pattern
    parts = re.split(r"<p><strong>Pytanie (\d+)\.</strong>", content)

    for i in range(1, len(parts), 2):
        number = int(parts[i])
        rest = parts[i + 1] if i + 1 < len(parts) else ""

        # Extract points and content
        match = re.match(r"\s*\((\d+) pkt\)\s*(.*?)</p>", rest, re.DOTALL)
        if match:
            points = int(match.group(1))
            q_content = re.sub(r"<[^>]+>", "", match.group(2)).strip()
        else:
            points = 1
            q_content = re.sub(r"<[^>]+>", "", rest).strip()

        # Check for options (closed question) — strip any nested HTML tags
        options = []
        option_matches = re.findall(r"<li>(.*?)</li>", rest, re.DOTALL)
        if option_matches:
            options = [re.sub(r"<[^>]+>", "", opt).strip() for opt in option_matches]

        q_type = "closed" if options else "open"

        questions.append({
            "number": number,
            "type": q_type,
            "content": q_content,
            "options": options,
            "correct_answer": "",
            "points": points,
        })

    return questions


def _parse_answer_key(answer_key_text: str) -> dict[int, str]:
    """Parse answer key text to a dict mapping question number → answer."""
    answers = {}
    for line in answer_key_text.splitlines():
        match = re.match(r"(\d+)\.\s*(.+)", line.strip())
        if match:
            answers[int(match.group(1))] = match.group(2).strip()
    return answers


_OPTION_LETTERS = {0: "a", 1: "b", 2: "c", 3: "d", 4: "e", 5: "f"}


def _add_free_form_content_to_docx(doc: DocxDocument, html_content: str, generation: Generation):
    """Convert basic HTML content to DOCX paragraphs for free-form types (worksheet/lesson_materials)."""
    from html.parser import HTMLParser

    class _HtmlDocxParser(HTMLParser):
        def __init__(self, doc: DocxDocument):
            super().__init__()
            self._doc = doc
            self._current_para = None
            self._current_run = None
            self._bold = False
            self._italic = False
            self._in_li = False
            self._heading_level = 0

        def _ensure_para(self):
            if self._current_para is None:
                self._current_para = self._doc.add_paragraph()

        def handle_starttag(self, tag, attrs):
            if tag in ("h1", "h2", "h3"):
                self._heading_level = int(tag[1])
                self._current_para = self._doc.add_heading(level=self._heading_level)
                self._current_run = None
            elif tag == "p":
                self._current_para = self._doc.add_paragraph()
                self._current_run = None
            elif tag in ("ul", "ol"):
                pass
            elif tag == "li":
                self._in_li = True
                self._current_para = self._doc.add_paragraph(style="List Bullet")
                self._current_run = None
            elif tag == "strong" or tag == "b":
                self._bold = True
            elif tag in ("em", "i"):
                self._italic = True
            elif tag == "br":
                self._ensure_para()
                self._current_para.add_run("\n")

        def handle_endtag(self, tag):
            if tag in ("h1", "h2", "h3", "p", "li"):
                self._current_para = None
                self._current_run = None
                self._in_li = False
                self._heading_level = 0
            elif tag in ("strong", "b"):
                self._bold = False
            elif tag in ("em", "i"):
                self._italic = False

        def handle_data(self, data):
            text = data
            if not text.strip() and not text:
                return
            self._ensure_para()
            run = self._current_para.add_run(text)
            run.bold = self._bold
            run.italic = self._italic

    parser = _HtmlDocxParser(doc)
    parser.feed(html_content)


def _generate_free_form_docx(doc: DocxDocument, prototype: Prototype, generation: Generation):
    """Generate DOCX content for worksheet or lesson_materials (HTML-based, not Q&A)."""
    content = prototype.edited_content or prototype.original_content or ""
    _add_free_form_content_to_docx(doc, content, generation)


def _strip_letter_prefix(text: str) -> str:
    """Strip option letter prefix like 'a) ' from answer text."""
    m = re.match(r"^[a-f]\)\s*(.*)", text.strip())
    return m.group(1).strip() if m else text.strip()


def _shuffle_variant(questions: list[dict], answer_map: dict[int, str]) -> tuple[list[dict], dict[int, str]]:
    """Shuffle questions and options within closed questions for a variant.

    Returns a new list of questions with updated numbering and the new answer map.
    The answer map is updated to reflect the new option letter after shuffling.
    """
    shuffled = copy.deepcopy(questions)
    # Work on a mutable copy of the map
    current_map = dict(answer_map)

    # Separate open and closed questions
    open_qs = [q for q in shuffled if q["type"] == "open"]
    closed_qs = [q for q in shuffled if q["type"] == "closed"]

    # Shuffle each group
    random.shuffle(open_qs)
    random.shuffle(closed_qs)

    # Shuffle options within closed questions and update answer letter
    for q in closed_qs:
        if q.get("options"):
            original_number = q["number"]
            correct_answer = current_map.get(original_number, "")
            correct_text = _strip_letter_prefix(correct_answer)

            random.shuffle(q["options"])

            # Find new letter position of the correct answer
            if correct_text:
                for idx, opt in enumerate(q["options"]):
                    if _strip_letter_prefix(opt).lower() == correct_text.lower():
                        new_letter = _OPTION_LETTERS.get(idx, str(idx + 1))
                        current_map[original_number] = f"{new_letter}) {correct_text}"
                        break

    # Merge back: closed first, then open
    merged = closed_qs + open_qs

    # Re-number questions and rebuild answer map with new sequential keys
    new_answer_map = {}
    for i, q in enumerate(merged, 1):
        old_number = q["number"]
        q["number"] = i
        if old_number in current_map:
            new_answer_map[i] = current_map[old_number]

    return merged, new_answer_map


def _add_questions_to_docx(doc: DocxDocument, questions: list[dict], group_label: str, generation: Generation):
    """Add a group of questions to the DOCX document."""
    # Title
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{generation.topic}")
    run.font.size = Pt(16)

    # Group label
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(group_label).font.size = Pt(10)

    # Date — left empty for student to fill in
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run("Data: ___________").font.size = Pt(10)

    # Separator
    doc.add_paragraph("_" * 60)

    # Questions
    for q in questions:
        q_para = doc.add_paragraph()
        q_para.add_run(f"Pytanie {q['number']}. ").bold = True
        q_para.add_run(f"({q.get('points', 1)} pkt) ")
        q_para.add_run(q.get("content", ""))

        if q["type"] == "closed" and q.get("options"):
            for opt in q["options"]:
                opt_para = doc.add_paragraph(style="List Bullet")
                opt_para.add_run(opt)

        if q["type"] == "open":
            # Add space for answer
            for _ in range(3):
                doc.add_paragraph("." * 80)


def _add_answer_key_to_docx(doc: DocxDocument, all_variants_answers: list[tuple[str, dict[int, str]]]):
    """Add answer key for all variants at the end of the document."""
    doc.add_page_break()

    heading = doc.add_heading("Klucz odpowiedzi", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for group_label, answer_map in all_variants_answers:
        doc.add_heading(group_label, level=2)
        for num in sorted(answer_map.keys()):
            doc.add_paragraph(f"{num}. {answer_map[num]}")


def generate_docx(db: DBSession, generation_id: str, user_id: str | None = None) -> Document:
    """Generate a DOCX file with variants from a prototype."""
    generation = db.query(Generation).filter(Generation.id == generation_id).first()
    if not generation:
        raise ValueError("Generation not found")

    prototype = db.query(Prototype).filter(Prototype.generation_id == generation_id).first()
    if not prototype:
        raise ValueError("Prototype not found")

    doc = DocxDocument()

    # Free-form types (worksheet / lesson_materials) — render HTML content, no variants
    if generation.content_type in TYPES_WITHOUT_QUESTIONS:
        _generate_free_form_docx(doc, prototype, generation)

        # Build path: content_type / education_level / class_level / subject
        subject = db.query(Subject).filter(Subject.id == generation.subject_id).first()
        subject_folder = re.sub(r'[<>:"/\\|?*]', '_', subject.name) if subject else generation_id
        content_type_folder = CONTENT_TYPE_FOLDER_NAMES.get(
            generation.content_type,
            re.sub(r'[<>:"/\\|?*]', '_', generation.content_type),
        )
        education_level_folder = EDUCATION_LEVEL_FOLDER_NAMES.get(
            generation.education_level,
            re.sub(r'[<>:"/\\|?*]', '_', str(generation.education_level).strip()) if generation.education_level else 'brak_poziomu',
        )
        class_level_folder = re.sub(r'[<>:"/\\|?*]', '_', str(generation.class_level).strip()) if generation.class_level else 'brak_klasy'
        docs_dir = Path(settings.DATA_DIR) / "documents" / content_type_folder / education_level_folder / class_level_folder / subject_folder
        docs_dir.mkdir(parents=True, exist_ok=True)

        filename = f"{generation.topic[:50].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        file_path = docs_dir / filename
        doc.save(str(file_path))

        document = Document(
            user_id=user_id or generation.user_id,
            generation_id=generation_id,
            filename=filename,
            file_path=str(file_path),
            variants_count=1,
        )
        db.add(document)
        generation.status = "finalized"
        generation.updated_at = datetime.now(timezone.utc).isoformat()
        db.commit()
        db.refresh(document)
        return document

    # Question-based types (test / quiz / exam / sprawdzian) — use variants
    # Prefer raw JSON (no HTML artifacts); fall back to HTML-parsed edited content
    questions = []
    answer_map = {}

    if prototype.raw_questions_json and not prototype.edited_content:
        try:
            raw_data = json.loads(prototype.raw_questions_json)
            questions = raw_data.get("questions", [])
            answer_map = {
                q["number"]: q.get("correct_answer", "")
                for q in questions
                if isinstance(q, dict) and q.get("number")
            }
        except (json.JSONDecodeError, TypeError):
            questions = []

    if not questions:
        content = prototype.edited_content or prototype.original_content
        questions = _parse_content_to_questions(content)
        answer_map = _parse_answer_key(prototype.answer_key)

    variants_count = generation.variants_count
    group_labels = [chr(65 + i) for i in range(variants_count)]  # A, B, C, ...
    all_variants_answers = []

    for i in range(variants_count):
        group_label = f"Grupa {group_labels[i]}"

        if i > 0:
            doc.add_page_break()

        if variants_count > 1:
            variant_questions, variant_answers = _shuffle_variant(questions, answer_map)
        else:
            variant_questions = questions
            variant_answers = answer_map

        _add_questions_to_docx(doc, variant_questions, group_label, generation)
        all_variants_answers.append((group_label, variant_answers))

    # Add answer key
    _add_answer_key_to_docx(doc, all_variants_answers)

    # Save file under hierarchical directory: content_type / education_level / class_level / subject
    subject = db.query(Subject).filter(Subject.id == generation.subject_id).first()
    subject_folder = re.sub(r'[<>:"/\\|?*]', '_', subject.name) if subject else generation_id
    content_type_folder = CONTENT_TYPE_FOLDER_NAMES.get(generation.content_type, re.sub(r'[<>:"/\\|?*]', '_', generation.content_type))
    education_level_folder = EDUCATION_LEVEL_FOLDER_NAMES.get(
        generation.education_level,
        re.sub(r'[<>:"/\\|?*]', '_', str(generation.education_level).strip()) if generation.education_level else 'brak_poziomu',
    )
    class_level_folder = re.sub(r'[<>:"/\\|?*]', '_', str(generation.class_level).strip()) if generation.class_level else 'brak_klasy'
    docs_dir = Path(settings.DATA_DIR) / "documents" / content_type_folder / education_level_folder / class_level_folder / subject_folder
    docs_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{generation.topic[:50].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    # Sanitize filename
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    file_path = docs_dir / filename
    doc.save(str(file_path))

    # Create DB record
    document = Document(
        user_id=user_id or generation.user_id,
        generation_id=generation_id,
        filename=filename,
        file_path=str(file_path),
        variants_count=variants_count,
    )
    db.add(document)

    generation.status = "finalized"
    generation.updated_at = datetime.now(timezone.utc).isoformat()
    db.commit()
    db.refresh(document)

    return document
